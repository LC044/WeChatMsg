import csv
import html
import os
import sys
import time
import traceback
from re import findall
from PyQt5.QtCore import pyqtSignal, QThread
from PyQt5.QtWidgets import QFileDialog
from docx.oxml.ns import qn

from . import msg_db, micro_msg_db
from .package_msg import PackageMsg
from ..DataBase import hard_link_db
from ..DataBase import media_msg_db
from ..log import logger
from ..person import MePC
from ..util import path
import shutil
from ..util.compress_content import parser_reply
from ..util.emoji import get_emoji, get_emoji_path, get_emoji_url
from ..util.image import get_image_path, get_image, get_image_abs_path
from ..util.file import get_file
import docx
from docx import shared
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT

os.makedirs('./data/聊天记录', exist_ok=True)

def set_global_font(doc, font_name):
    # 创建一个新样式
    style = doc.styles['Normal']

    # 设置字体名称
    style.font.name = font_name
    # 遍历文档中的所有段落，将样式应用到每个段落
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
def makedirs(path):
    os.makedirs(path, exist_ok=True)
    os.makedirs(os.path.join(path, 'image'), exist_ok=True)
    os.makedirs(os.path.join(path, 'emoji'), exist_ok=True)
    os.makedirs(os.path.join(path, 'video'), exist_ok=True)
    os.makedirs(os.path.join(path, 'voice'), exist_ok=True)
    os.makedirs(os.path.join(path, 'file'), exist_ok=True)
    os.makedirs(os.path.join(path, 'avatar'), exist_ok=True)
    file = './app/resources/data/file.png'
    if not os.path.exists(file):
        resource_dir = getattr(sys, '_MEIPASS', os.path.abspath(os.path.dirname(__file__)))
        file = os.path.join(resource_dir, 'app', 'resources', 'data', 'file.png')
    shutil.copy(file, path + '/file/file.png')


def escape_js_and_html(input_str):
    # 转义HTML特殊字符
    html_escaped = html.escape(input_str, quote=False)

    # 手动处理JavaScript转义字符
    js_escaped = (
        html_escaped
        .replace("\\", "\\\\")
        .replace("'", r"\'")
        .replace('"', r'\"')
        .replace("\n", r'\n')
        .replace("\r", r'\r')
        .replace("\t", r'\t')
    )

    return js_escaped


class Output(QThread):
    """
    发送信息线程
    """
    progressSignal = pyqtSignal(int)
    rangeSignal = pyqtSignal(int)
    okSignal = pyqtSignal(int)
    i = 1
    CSV = 0
    DOCX = 1
    HTML = 2
    CSV_ALL = 3
    CONTACT_CSV = 4
    TXT = 5

    def __init__(self, contact, type_=DOCX, message_types={}, parent=None):
        super().__init__(parent)
        self.Child0 = None
        self.last_timestamp = 0
        self.message_types = message_types
        self.sec = 2  # 默认1000秒
        self.contact = contact
        self.ta_username = contact.wxid if contact else ''
        self.msg_id = 0
        self.output_type = type_
        self.total_num = 1
        self.num = 0

    def progress(self, value):
        self.progressSignal.emit(value)

    def to_csv_all(self):
        origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/"
        os.makedirs(origin_docx_path, exist_ok=True)
        filename = QFileDialog.getSaveFileName(None, "save file", os.path.join(os.getcwd(), 'messages.csv'),
                                               "csv files (*.csv);;all files(*.*)")
        if not filename[0]:
            return
        filename = filename[0]
        # columns = ["用户名", "消息内容", "发送时间", "发送状态", "消息类型", "isSend", "msgId"]
        columns = ['localId', 'TalkerId', 'Type', 'SubType',
                   'IsSender', 'CreateTime', 'Status', 'StrContent',
                   'StrTime', 'Remark', 'NickName', 'Sender']

        packagemsg = PackageMsg()
        messages = packagemsg.get_package_message_all()
        # 写入CSV文件
        with open(filename, mode='w', newline='', encoding='utf-8-sig') as file:
            writer = csv.writer(file)
            writer.writerow(columns)
            # 写入数据
            writer.writerows(messages)
        self.okSignal.emit(1)

    def contact_to_csv(self):
        filename = QFileDialog.getSaveFileName(None, "save file", os.path.join(os.getcwd(), 'contacts.csv'),
                                               "csv files (*.csv);;all files(*.*)")
        if not filename[0]:
            return
        filename = filename[0]
        # columns = ["用户名", "消息内容", "发送时间", "发送状态", "消息类型", "isSend", "msgId"]
        columns = ['UserName', 'Alias', 'Type', 'Remark', 'NickName', 'PYInitial', 'RemarkPYInitial', 'smallHeadImgUrl',
                   'bigHeadImgUrl']
        contacts = micro_msg_db.get_contact()
        # 写入CSV文件
        with open(filename, mode='w', newline='', encoding='utf-8-sig') as file:
            writer = csv.writer(file)
            writer.writerow(columns)
            # 写入数据
            writer.writerows(contacts)
        self.okSignal.emit(1)

    def run(self):
        if self.output_type == self.DOCX:
            self.Child = ChildThread(self.contact, type_=self.output_type, message_types=self.message_types)
            self.Child.progressSignal.connect(self.progress)
            self.Child.rangeSignal.connect(self.rangeSignal)
            self.Child.okSignal.connect(self.okSignal)
            self.Child.start()
        elif self.output_type == self.CSV_ALL:
            self.to_csv_all()
        elif self.output_type == self.CONTACT_CSV:
            self.contact_to_csv()
        elif self.output_type == self.CSV or self.output_type == self.TXT or self.output_type == self.DOCX:
            self.Child = ChildThread(self.contact, type_=self.output_type, message_types=self.message_types)
            self.Child.progressSignal.connect(self.progress)
            self.Child.rangeSignal.connect(self.rangeSignal)
            self.Child.okSignal.connect(self.okSignal)
            self.Child.start()
        elif self.output_type == self.HTML:
            self.Child = ChildThread(self.contact, type_=self.output_type, message_types=self.message_types)
            self.Child.progressSignal.connect(self.progressSignal)
            self.Child.rangeSignal.connect(self.rangeSignal)
            self.Child.okSignal.connect(self.count_finish_num)
            self.Child.start()
            if self.message_types.get(34):
                # 语音消息单独的线程
                self.total_num += 1
                self.output_media = OutputMedia(self.contact)
                self.output_media.okSingal.connect(self.count_finish_num)
                self.output_media.progressSignal.connect(self.progressSignal)
                self.output_media.start()

            if self.message_types.get(47):
                # emoji消息单独的线程
                self.total_num += 1
                self.output_emoji = OutputEmoji(self.contact)
                self.output_emoji.okSingal.connect(self.count_finish_num)
                self.output_emoji.progressSignal.connect(self.progressSignal)
                self.output_emoji.start()
            if self.message_types.get(3):
                # 图片消息单独的线程
                self.total_num += 1
                self.output_image = OutputImage(self.contact)
                self.output_image.okSingal.connect(self.count_finish_num)
                self.output_image.progressSignal.connect(self.progressSignal)
                self.output_image.start()

    def count_finish_num(self, num):
        self.num += 1
        if self.num == self.total_num:
            self.okSignal.emit(1)

    def cancel(self):
        self.requestInterruption()


def modify_audio_metadata(audiofile, new_artist):  # 修改音频元数据中的“创作者”标签
    return
    audiofile = load(audiofile)

    # 检查文件是否有标签
    if audiofile.tag is None:
        audiofile.initTag()

    # 修改艺术家名称
    audiofile.tag.artist = new_artist
    audiofile.tag.save()


class ChildThread(QThread):
    """
        子线程，用于导出部分聊天记录
    """
    progressSignal = pyqtSignal(int)
    rangeSignal = pyqtSignal(int)
    okSignal = pyqtSignal(int)
    i = 1
    CSV = 0
    DOCX = 1
    HTML = 2

    def __init__(self, contact, type_=DOCX, message_types={}, parent=None):
        super().__init__(parent)
        self.contact = contact
        self.message_types = message_types
        self.last_timestamp = 0
        self.sec = 2  # 默认1000秒
        self.msg_id = 0
        self.output_type = type_

    def is_5_min(self, timestamp):
        if abs(timestamp - self.last_timestamp) > 300:
            self.last_timestamp = timestamp
            return True
        return False

    def text(self, doc, message):
        type_ = message[2]
        str_content = message[7]
        str_time = message[8]
        is_send = message[4]
        timestamp = message[5]
        is_chatroom = 1 if self.contact.is_chatroom else 0
        if is_chatroom:
            avatar = f"./avatar/{message[12].wxid}.png"
        else:
            avatar = f"./avatar/{MePC().wxid if is_send else self.contact.wxid}.png"
        if is_chatroom:
            if is_send:
                displayname = MePC().name
            else:
                displayname = message[12].remark
        else:
            displayname = MePC().name if is_send else self.contact.remark
        displayname = escape_js_and_html(displayname)
        if self.output_type == Output.HTML:
            str_content = escape_js_and_html(str_content)
            doc.write(
                f'''{{ type:{1}, text: '{str_content}',is_send:{is_send},avatar_path:'{avatar}',timestamp:{timestamp},is_chatroom:{is_chatroom},displayname:'{displayname}'}},'''
            )
        elif self.output_type == Output.TXT:
            name = displayname
            doc.write(
                f'''{str_time} {name}\n{str_content}\n\n'''
            )
        elif self.output_type == Output.DOCX:
            origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
            if is_chatroom:
                avatar = f"{origin_docx_path}/avatar/{message[12].wxid}.png"
            else:
                avatar = f"{origin_docx_path}/avatar/{MePC().wxid if is_send else self.contact.wxid}.png"
            content_cell = self.create_table(doc, is_send, avatar)
            content_cell.paragraphs[0].add_run(str_content)
            content_cell.paragraphs[0].font_size = shared.Inches(0.5)
            if is_send:
                p = content_cell.paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            doc.add_paragraph()

    def image(self, doc, message):
        origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
        type_ = message[2]
        str_content = message[7]
        str_time = message[8]
        is_send = message[4]
        BytesExtra = message[10]
        timestamp = message[5]
        is_chatroom = 1 if self.contact.is_chatroom else 0
        if is_chatroom:
            avatar = f"./avatar/{message[12].wxid}.png"
        else:
            avatar = f"./avatar/{MePC().wxid if is_send else self.contact.wxid}.png"
        if is_chatroom:
            if is_send:
                displayname = MePC().name
            else:
                displayname = message[12].remark
        else:
            displayname = MePC().name if is_send else self.contact.remark
        displayname = escape_js_and_html(displayname)
        if self.output_type == Output.HTML:
            str_content = escape_js_and_html(str_content)
            image_path = hard_link_db.get_image(str_content, BytesExtra, thumb=False)
            if not os.path.exists(os.path.join(MePC().wx_dir, image_path)):
                image_thumb_path = hard_link_db.get_image(str_content, BytesExtra, thumb=True)
                if not os.path.exists(os.path.join(MePC().wx_dir, image_thumb_path)):
                    return
                image_path = image_thumb_path
            image_path = get_image_path(image_path, base_path=f'/data/聊天记录/{self.contact.remark}/image')
            image_path = image_path.replace('/', '\\')
            image_path = image_path.replace('\\', '/')
            doc.write(
                f'''{{ type:{type_}, text: '{image_path}',is_send:{is_send},avatar_path:'{avatar}',timestamp:{timestamp},is_chatroom:{is_chatroom},displayname:'{displayname}'}},'''
            )
        elif self.output_type == Output.TXT:
            name = displayname
            doc.write(
                f'''{str_time} {name}\n[图片]\n\n'''
            )
        elif self.output_type == Output.DOCX:
            origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
            if is_chatroom:
                avatar = f"{origin_docx_path}/avatar/{message[12].wxid}.png"
            else:
                avatar = f"{origin_docx_path}/avatar/{MePC().wxid if is_send else self.contact.wxid}.png"
            content = self.create_table(doc, is_send,avatar)
            run = content.paragraphs[0].add_run()
            str_content = escape_js_and_html(str_content)
            image_path = hard_link_db.get_image(str_content, BytesExtra, thumb=True)
            if not os.path.exists(os.path.join(MePC().wx_dir, image_path)):
                image_thumb_path = hard_link_db.get_image(str_content, BytesExtra, thumb=False)
                if not os.path.exists(os.path.join(MePC().wx_dir, image_thumb_path)):
                    return
                image_path = image_thumb_path
            image_path = get_image_abs_path(image_path, base_path=f'/data/聊天记录/{self.contact.remark}/image')

            try:
                run.add_picture(image_path, height=shared.Inches(2))
                doc.add_paragraph()
            except Exception:
                print("Error!image")

    def audio(self, doc, message):
        origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
        str_content = message[7]
        str_time = message[8]
        is_send = message[4]
        msgSvrId = message[9]
        timestamp = message[5]
        is_chatroom = 1 if self.contact.is_chatroom else 0
        if is_chatroom:
            avatar = f"./avatar/{message[12].wxid}.png"
        else:
            avatar = f"./avatar/{MePC().wxid if is_send else self.contact.wxid}.png"
        if is_chatroom:
            if is_send:
                displayname = MePC().name
            else:
                displayname = message[12].remark
        else:
            displayname = MePC().name if is_send else self.contact.remark
        displayname = escape_js_and_html(displayname)
        if self.output_type == Output.HTML:
            try:
                audio_path = media_msg_db.get_audio_path(msgSvrId, output_path=origin_docx_path + "/voice")
                audio_path = audio_path.replace('/', '\\')
                audio_path = audio_path.replace('\\', '/')
                audio_path = "./voice/" + os.path.basename(audio_path)
                voice_to_text = escape_js_and_html(media_msg_db.get_audio_text(str_content))
            except:
                logger.error(traceback.format_exc())
                return
            doc.write(
                f'''{{ type:34, text:'{audio_path}',is_send:{is_send},avatar_path:'{avatar}',voice_to_text:'{voice_to_text}',timestamp:{timestamp},is_chatroom:{is_chatroom},displayname:'{displayname}'}},'''
            )
        if self.output_type == Output.TXT:
            name = displayname
            doc.write(
                f'''{str_time} {name}\n[语音]\n\n'''
            )

    def emoji(self, doc, message):
        origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
        str_content = message[7]
        str_time = message[8]
        is_send = message[4]
        timestamp = message[5]
        is_chatroom = 1 if self.contact.is_chatroom else 0
        if is_chatroom:
            avatar = f"./avatar/{message[12].wxid}.png"
        else:
            avatar = f"./avatar/{MePC().wxid if is_send else self.contact.wxid}.png"
        if is_chatroom:
            if is_send:
                displayname = MePC().name
            else:
                displayname = message[12].remark
        else:
            displayname = MePC().name if is_send else self.contact.remark
        displayname = escape_js_and_html(displayname)
        if self.output_type == Output.HTML:
            # emoji_path = get_emoji_path(str_content, thumb=True, output_path=origin_docx_path + '/emoji')
            # emoji_path = './emoji/' + os.path.basename(emoji_path)
            emoji_path = get_emoji_url(str_content, thumb=True)
            doc.write(
                f'''{{ type:{3}, text: '{emoji_path}',is_send:{is_send},avatar_path:'{avatar}',timestamp:{timestamp},is_chatroom:{is_chatroom},displayname:'{displayname}'}},'''
            )
        elif self.output_type == Output.TXT:
            name = displayname
            doc.write(
                f'''{str_time} {name}\n[表情包]\n\n'''
            )
        elif self.output_type == Output.DOCX:
            origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
            if is_chatroom:
                avatar = f"{origin_docx_path}/avatar/{message[12].wxid}.png"
            else:
                avatar = f"{origin_docx_path}/avatar/{MePC().wxid if is_send else self.contact.wxid}.png"
            content_cell = self.create_table(doc, is_send, avatar)
            content_cell.paragraphs[0].add_run('【表情包】')
            content_cell.paragraphs[0].font_size = shared.Inches(0.5)
            if is_send:
                p = content_cell.paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            doc.add_paragraph()

    def wx_file(self, doc, isSend, content, status):
        return

    def file(self, doc, message):
        origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
        bytesExtra = message[10]
        str_time = message[8]
        is_send = message[4]
        timestamp = message[5]
        is_chatroom = 1 if self.contact.is_chatroom else 0
        if is_chatroom:
            avatar = f"./avatar/{message[12].wxid}.png"
        else:
            avatar = f"./avatar/{MePC().wxid if is_send else self.contact.wxid}.png"
        if is_chatroom:
            if is_send:
                displayname = MePC().name
            else:
                displayname = message[12].remark
        else:
            displayname = MePC().name if is_send else self.contact.remark
        displayname = escape_js_and_html(displayname)
        if self.output_type == Output.HTML:
            link = get_file(bytesExtra, thumb=True, output_path=origin_docx_path + '/file')
            file_name = ''
            file_path = './file/file.png'
            if link != "":
                file_name = os.path.basename(link)
                link = './file/' + file_name
            doc.write(
                f'''{{ type:49, text: '{file_path}',is_send:{is_send},avatar_path:'{avatar}',timestamp:{timestamp},is_chatroom:{is_chatroom},displayname:'{displayname}',link: '{link}',sub_type:6,file_name: '{file_name}'}},'''
            )
        elif self.output_type == Output.TXT:
            if is_chatroom:
                if is_send:
                    displayname = MePC().name
                else:
                    displayname = message[12].remark
            else:
                displayname = MePC().name if is_send else self.contact.remark
            doc.write(
                f'''{str_time} {displayname}\n[文件]\n\n'''
            )
        elif self.output_type == Output.DOCX:
            origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
            if is_chatroom:
                avatar = f"{origin_docx_path}/avatar/{message[12].wxid}.png"
            else:
                avatar = f"{origin_docx_path}/avatar/{MePC().wxid if is_send else self.contact.wxid}.png"
            content_cell = self.create_table(doc, is_send, avatar)
            content_cell.paragraphs[0].add_run('【文件】')
            content_cell.paragraphs[0].font_size = shared.Inches(0.5)
            if is_send:
                p = content_cell.paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            doc.add_paragraph()

    def refermsg(self, doc, message):
        """
        处理回复消息
        @param doc:
        @param message:
        @return:
        """
        str_time = message[8]
        is_send = message[4]
        content = parser_reply(message[11])
        refer_msg = content.get('refer')
        timestamp = message[5]
        is_chatroom = 1 if self.contact.is_chatroom else 0
        if is_chatroom:
            avatar = f"./avatar/{message[12].wxid}.png"
        else:
            avatar = f"./avatar/{MePC().wxid if is_send else self.contact.wxid}.png"
        if is_chatroom:
            if is_send:
                displayname = MePC().name
            else:
                displayname = message[12].remark
        else:
            displayname = MePC().name if is_send else self.contact.remark
        displayname = escape_js_and_html(displayname)
        if self.output_type == Output.HTML:
            contentText = escape_js_and_html(content.get('title'))
            if refer_msg:
                referText = f"{escape_js_and_html(refer_msg.get('displayname'))}：{escape_js_and_html(refer_msg.get('content'))}"
                doc.write(
                    f'''{{ type:49, text: '{contentText}',is_send:{is_send},sub_type:{content.get('type')},refer_text: '{referText}',avatar_path:'{avatar}',timestamp:{timestamp},is_chatroom:{is_chatroom},displayname:'{displayname}'}},'''
                )
            else:
                doc.write(
                    f'''{{ type:49, text: '{contentText}',is_send:{is_send},sub_type:{content.get('type')},avatar_path:'{avatar}',timestamp:{timestamp},is_chatroom:{is_chatroom},displayname:'{displayname}'}},'''
                )
        elif self.output_type == Output.TXT:
            name = displayname
            if refer_msg:
                doc.write(
                    f'''{str_time} {name}\n{content.get('title')}\n引用:{refer_msg.get('displayname')}:{refer_msg.get('content')}\n\n'''
                )
            else:
                doc.write(
                    f'''{str_time} {name}\n{content.get('title')}\n引用:未知\n\n'''
                )
        elif self.output_type == Output.DOCX:
            origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
            if is_chatroom:
                avatar = f"{origin_docx_path}/avatar/{message[12].wxid}.png"
            else:
                avatar = f"{origin_docx_path}/avatar/{MePC().wxid if is_send else self.contact.wxid}.png"
            content_cell = self.create_table(doc, is_send, avatar)

            content_cell.paragraphs[0].add_run(content.get('title'))
            content_cell.paragraphs[0].font_size = shared.Inches(0.5)
            reply_p = content_cell.add_paragraph()
            reply_content = f"{refer_msg.get('displayname')}:{refer_msg.get('content')}" if refer_msg else '未知引用'
            run = content_cell.paragraphs[1].add_run(reply_content)
            '''设置被回复内容格式'''
            run.font.color.rgb = shared.RGBColor(121, 121, 121)
            run.font_size = shared.Inches(0.3)
            run.font.highlight_color = WD_COLOR_INDEX.GRAY_25

            if is_send:
                p = content_cell.paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                reply_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            doc.add_paragraph()
    def system_msg(self, doc, message):
        str_content = message[7]
        is_send = message[4]
        str_time = message[8]
        timestamp = message[5]
        is_chatroom = 1 if self.contact.is_chatroom else 0
        if is_chatroom:
            if is_send:
                displayname = MePC().name
            else:
                displayname = message[12].remark
        else:
            displayname = MePC().name if is_send else self.contact.remark
        str_content = str_content.replace('<![CDATA[', "").replace(
            ' <a href="weixin://revoke_edit_click">重新编辑</a>]]>', "")
        res = findall('(</{0,1}(img|revo|_wc_cus|a).*?>)', str_content)
        for xmlstr, b in res:
            str_content = str_content.replace(xmlstr, "")
        if self.output_type == Output.HTML:
            str_content = escape_js_and_html(str_content)
            doc.write(
                f'''{{ type:0, text: '{str_content}',is_send:{is_send},avatar_path:'',timestamp:{timestamp},is_chatroom:{is_chatroom},displayname:''}},'''
            )
        elif self.output_type == Output.TXT:

            doc.write(
                f'''{str_time} {displayname}\n{str_content}\n\n'''
            )
        elif self.output_type == Output.DOCX:
            doc.add_paragraph(str_content).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def video(self, doc, message):
        origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
        type_ = message[2]
        str_content = message[7]
        str_time = message[8]
        is_send = message[4]
        BytesExtra = message[10]
        timestamp = message[5]
        is_chatroom = 1 if self.contact.is_chatroom else 0
        if is_chatroom:
            avatar = f"./avatar/{message[12].wxid}.png"
        else:
            avatar = f"./avatar/{MePC().wxid if is_send else self.contact.wxid}.png"
        if is_chatroom:
            if is_send:
                displayname = MePC().name
            else:
                displayname = message[12].remark
        else:
            displayname = MePC().name if is_send else self.contact.remark
        displayname = escape_js_and_html(displayname)
        if self.output_type == Output.HTML:
            video_path = hard_link_db.get_video(str_content, BytesExtra, thumb=False)
            image_path = hard_link_db.get_video(str_content, BytesExtra, thumb=True)
            if video_path is None and image_path is not None:
                image_path = path.get_relative_path(image_path, base_path=f'/data/聊天记录/{self.contact.remark}/image')
                try:
                    # todo 网络图片问题
                    print(origin_docx_path + image_path[1:])
                    os.utime(origin_docx_path + image_path[1:], (timestamp, timestamp))
                    image_path = image_path.replace('\\', '/')
                    # print(f"tohtml:---{image_path}")

                    doc.write(
                        f'''{{ type:3, text: '{image_path}',is_send:{is_send},avatar_path:'{avatar}',timestamp:{timestamp},is_chatroom:{is_chatroom},displayname:'{displayname}'}},'''
                    )
                except:
                    doc.write(
                        f'''{{ type:1, text: '视频丢失',is_send:{is_send},avatar_path:'{avatar}',timestamp:{timestamp},is_chatroom:{is_chatroom},displayname:'{displayname}'}},'''
                    )
                return
            if video_path is None and image_path is None:
                return
            video_path = f'{MePC().wx_dir}/{video_path}'
            if os.path.exists(video_path):
                new_path = origin_docx_path + '/video/' + os.path.basename(video_path)
                if not os.path.exists(new_path):
                    shutil.copy(video_path, os.path.join(origin_docx_path, 'video'))
                os.utime(new_path, (timestamp, timestamp))
                video_path = f'./video/{os.path.basename(video_path)}'
            video_path = video_path.replace('\\', '/')
            doc.write(
                f'''{{ type:{type_}, text: '{video_path}',is_send:{is_send},avatar_path:'{avatar}',timestamp:{timestamp},is_chatroom:{is_chatroom},displayname:'{displayname}'}},'''
            )
        elif self.output_type == Output.TXT:
            if is_chatroom:
                if is_send:
                    displayname = MePC().name
                else:
                    displayname = message[12].remark
            else:
                displayname = MePC().name if is_send else self.contact.remark
            doc.write(
                f'''{str_time} {displayname}\n[视频]\n\n'''
            )
        elif self.output_type == Output.DOCX:
            origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
            if is_chatroom:
                avatar = f"{origin_docx_path}/avatar/{message[12].wxid}.png"
            else:
                avatar = f"{origin_docx_path}/avatar/{MePC().wxid if is_send else self.contact.wxid}.png"
            content_cell = self.create_table(doc, is_send, avatar)
            content_cell.paragraphs[0].add_run('【视频】')
            content_cell.paragraphs[0].font_size = shared.Inches(0.5)
            if is_send:
                p = content_cell.paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            doc.add_paragraph()
    def create_table(self, doc, is_send, avatar_path):
        '''
        #! 创建一个1*2表格
        #! isSend = 1 (0,0)存聊天内容，(0,1)存头像
        #! isSend = 0 (0,0)存头像，(0,1)存聊天内容
        #! 返回聊天内容的坐标
        '''
        table = doc.add_table(rows=1, cols=2, style='Normal Table')
        table.cell(0, 1).height = shared.Inches(0.5)
        table.cell(0, 0).height = shared.Inches(0.5)
        if is_send:
            '''表格右对齐'''
            table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            avatar = table.cell(0, 1).paragraphs[0].add_run()
            '''插入头像，设置头像宽度'''
            avatar.add_picture(avatar_path, width=shared.Inches(0.5))
            '''设置单元格宽度跟头像一致'''
            table.cell(0, 1).width = shared.Inches(0.5)
            content_cell = table.cell(0, 0)
            '''聊天内容右对齐'''
            content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            avatar = table.cell(0, 0).paragraphs[0].add_run()
            avatar.add_picture(avatar_path, width=shared.Inches(0.5))
            '''设置单元格宽度'''
            table.cell(0, 0).width = shared.Inches(0.5)
            content_cell = table.cell(0, 1)
        '''聊天内容垂直居中对齐'''
        content_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        return content_cell

    def to_csv(self):
        origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
        os.makedirs(origin_docx_path, exist_ok=True)
        filename = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}/{self.contact.remark}_utf8.csv"
        # columns = ["用户名", "消息内容", "发送时间", "发送状态", "消息类型", "isSend", "msgId"]
        columns = ['localId', 'TalkerId', 'Type', 'SubType',
                   'IsSender', 'CreateTime', 'Status', 'StrContent',
                   'StrTime']
        messages = msg_db.get_messages(self.contact.wxid)
        # 写入CSV文件
        with open(filename, mode='w', newline='', encoding='utf-8-sig') as file:
            writer = csv.writer(file)
            writer.writerow(columns)
            # 写入数据
            writer.writerows(messages)
        self.okSignal.emit('ok')

    def to_html_(self):
        origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
        makedirs(origin_docx_path)

        if self.contact.is_chatroom:
            packagemsg = PackageMsg()
            messages = packagemsg.get_package_message_by_wxid(self.contact.wxid)
        else:
            messages = msg_db.get_messages(self.contact.wxid)
        filename = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}/{self.contact.remark}.html"
        file_path = './app/resources/data/template.html'
        if not os.path.exists(file_path):
            resource_dir = getattr(sys, '_MEIPASS', os.path.abspath(os.path.dirname(__file__)))
            file_path = os.path.join(resource_dir, 'app', 'resources', 'data', 'template.html')

        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
            html_head, html_end = content.split('/*注意看这是分割线*/')
        f = open(filename, 'w', encoding='utf-8')
        f.write(html_head.replace("<title>Chat Records</title>", f"<title>{self.contact.remark}</title>"))
        MePC().avatar.save(os.path.join(f"{origin_docx_path}/avatar/{MePC().wxid}.png"))
        if self.contact.is_chatroom:
            for message in messages:
                if message[4]:  # is_send
                    continue
                try:
                    chatroom_avatar_path = f"{origin_docx_path}/avatar/{message[12].wxid}.png"
                    if not os.path.exists(chatroom_avatar_path):
                        message[12].avatar.save(chatroom_avatar_path)
                except:
                    print(message)
                    pass
        else:
            self.contact.avatar.save(os.path.join(f"{origin_docx_path}/avatar/{self.contact.wxid}.png"))
        self.rangeSignal.emit(len(messages))
        for index, message in enumerate(messages):
            type_ = message[2]
            sub_type = message[3]
            timestamp = message[5]
            if (type_ == 3 and self.message_types.get(3)) or (type_ == 34 and self.message_types.get(34)) or (
                    type_ == 47 and self.message_types.get(47)):
                pass
            else:
                self.progressSignal.emit(1)

            if self.is_5_min(timestamp):
                str_time = message[8]
                f.write(
                    f'''{{ type:0, text: '{str_time}',is_send:0,avatar_path:'',timestamp:{timestamp}}},'''
                )
            if type_ == 1 and self.message_types.get(type_):
                self.text(f, message)
            elif type_ == 3 and self.message_types.get(type_):
                self.image(f, message)
            elif type_ == 34 and self.message_types.get(type_):
                self.audio(f, message)
            elif type_ == 43 and self.message_types.get(type_):
                self.video(f, message)
            elif type_ == 47 and self.message_types.get(type_):
                self.emoji(f, message)
            elif type_ == 10000 and self.message_types.get(type_):
                self.system_msg(f, message)
            elif type_ == 49 and sub_type == 57 and self.message_types.get(1):
                self.refermsg(f, message)
            elif type_ == 49 and sub_type == 6 and self.message_types.get(4906):
                self.file(f, message)
        f.write(html_end)
        f.close()
        self.okSignal.emit(1)

    def to_txt(self):
        origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
        os.makedirs(origin_docx_path, exist_ok=True)
        filename = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}/{self.contact.remark}.txt"
        if self.contact.is_chatroom:
            packagemsg = PackageMsg()
            messages = packagemsg.get_package_message_by_wxid(self.contact.wxid)
        else:
            messages = msg_db.get_messages(self.contact.wxid)
        total_steps = len(messages)
        with open(filename, mode='w', newline='', encoding='utf-8') as f:
            for index, message in enumerate(messages):
                type_ = message[2]
                sub_type = message[3]
                self.progressSignal.emit(int((index + 1) / total_steps * 100))
                if type_ == 1 and self.message_types.get(type_):
                    self.text(f, message)
                elif type_ == 3 and self.message_types.get(type_):
                    self.image(f, message)
                elif type_ == 34 and self.message_types.get(type_):
                    self.audio(f, message)
                elif type_ == 43 and self.message_types.get(type_):
                    self.video(f, message)
                elif type_ == 47 and self.message_types.get(type_):
                    self.emoji(f, message)
                elif type_ == 10000 and self.message_types.get(type_):
                    self.system_msg(f, message)
                elif type_ == 49 and sub_type == 57:
                    self.refermsg(f, message)
        self.okSignal.emit(1)

    def to_docx(self):
        print('导出docx')
        origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
        filename = os.path.join(origin_docx_path,f"{self.contact.remark}.docx")
        makedirs(origin_docx_path)
        doc = docx.Document()
        doc.styles['Normal'].font.name = u'Cambria'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        if self.contact.is_chatroom:
            packagemsg = PackageMsg()
            messages = packagemsg.get_package_message_by_wxid(self.contact.wxid)
        else:
            messages = msg_db.get_messages(self.contact.wxid)

        MePC().avatar.save(os.path.join(f"{origin_docx_path}/avatar/{MePC().wxid}.png"))
        if self.contact.is_chatroom:
            for message in messages:
                if message[4]:  # is_send
                    continue
                try:
                    chatroom_avatar_path = f"{origin_docx_path}/avatar/{message[12].wxid}.png"
                    if not os.path.exists(chatroom_avatar_path):
                        message[12].avatar.save(chatroom_avatar_path)
                except:
                    print(message)
                    pass
        else:
            self.contact.avatar.save(os.path.join(f"{origin_docx_path}/avatar/{self.contact.wxid}.png"))
        self.rangeSignal.emit(len(messages))
        for index, message in enumerate(messages):
            type_ = message[2]
            sub_type = message[3]
            timestamp = message[5]
            self.progressSignal.emit(1)
            if self.is_5_min(timestamp):
                str_time = message[8]
                doc.add_paragraph(str_time).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if type_ == 1 and self.message_types.get(type_):
                self.text(doc, message)
            elif type_ == 3 and self.message_types.get(type_):
                self.image(doc, message)
            elif type_ == 34 and self.message_types.get(type_):
                self.audio(doc, message)
            elif type_ == 43 and self.message_types.get(type_):
                self.video(doc, message)
            elif type_ == 47 and self.message_types.get(type_):
                self.emoji(doc, message)
            elif type_ == 10000 and self.message_types.get(type_):
                self.system_msg(doc, message)
            elif type_ == 49 and sub_type == 57 and self.message_types.get(1):
                self.refermsg(doc, message)
            elif type_ == 49 and sub_type == 6 and self.message_types.get(4906):
                self.file(doc, message)
        try:
            doc.save(filename)
        except PermissionError:
            filename = filename[:-5]+f'{time.time()}'+'.docx'
            doc.save(filename)
        self.okSignal.emit(1)

    def run(self):
        if self.output_type == Output.DOCX:
            self.to_docx()
        elif self.output_type == Output.CSV:
            self.to_csv()
        elif self.output_type == Output.HTML:
            self.to_html_()
        elif self.output_type == Output.CSV_ALL:
            self.to_csv_all()
        elif self.output_type == Output.TXT:
            self.to_txt()

    def cancel(self):
        self.requestInterruption()


class OutputMedia(QThread):
    okSingal = pyqtSignal(int)
    progressSignal = pyqtSignal(int)

    def __init__(self, contact):
        super().__init__()
        self.contact = contact

    def run(self):
        origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
        messages = msg_db.get_messages_by_type(self.contact.wxid, 34)
        for message in messages:
            is_send = message[4]
            msgSvrId = message[9]
            try:
                audio_path = media_msg_db.get_audio(msgSvrId, output_path=origin_docx_path + "/voice")
            except:
                logger.error(traceback.format_exc())
            finally:
                self.progressSignal.emit(1)
        self.okSingal.emit(34)


class OutputEmoji(QThread):
    okSingal = pyqtSignal(int)
    progressSignal = pyqtSignal(int)

    def __init__(self, contact):
        super().__init__()
        self.contact = contact

    def run(self):
        origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
        messages = msg_db.get_messages_by_type(self.contact.wxid, 47)
        for message in messages:
            str_content = message[7]
            try:
                pass
                # emoji_path = get_emoji(str_content, thumb=True, output_path=origin_docx_path + '/emoji')
            except:
                logger.error(traceback.format_exc())
            finally:
                self.progressSignal.emit(1)
        self.okSingal.emit(47)


class OutputImage(QThread):
    okSingal = pyqtSignal(int)
    progressSignal = pyqtSignal(int)

    def __init__(self, contact):
        super().__init__()
        self.contact = contact
        self.child_thread_num = 2
        self.child_threads = [0] * (self.child_thread_num + 1)
        self.num = 0

    def count1(self, num):
        self.num += 1
        print('图片导出完成一个')
        if self.num == self.child_thread_num:
            self.okSingal.emit(47)
            print('图片导出完成')

    def run(self):
        origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
        messages = msg_db.get_messages_by_type(self.contact.wxid, 3)
        for message in messages:
            str_content = message[7]
            BytesExtra = message[10]
            timestamp = message[5]
            try:
                image_path = hard_link_db.get_image(str_content, BytesExtra, thumb=False)
                if not os.path.exists(os.path.join(MePC().wx_dir, image_path)):
                    image_thumb_path = hard_link_db.get_image(str_content, BytesExtra, thumb=True)
                    if not os.path.exists(os.path.join(MePC().wx_dir, image_thumb_path)):
                        continue
                    image_path = image_thumb_path
                image_path = get_image(image_path, base_path=f'/data/聊天记录/{self.contact.remark}/image')
                try:
                    os.utime(origin_docx_path + image_path[1:], (timestamp, timestamp))
                except:
                    pass
            except:
                logger.error(traceback.format_exc())
            finally:
                self.progressSignal.emit(1)
        self.okSingal.emit(47)
        # sublist_length = len(messages) // self.child_thread_num
        # index = 0
        # for i in range(0, len(messages), sublist_length):
        #     child_messages = messages[i:i + sublist_length]
        #     self.child_threads[index] = OutputImageChild(self.contact, child_messages)
        #     self.child_threads[index].okSingal.connect(self.count1)
        #     self.child_threads[index].progressSignal.connect(self.progressSignal)
        #     self.child_threads[index].start()
        #     print('开启一个新线程')
        #     index += 1


class OutputImageChild(QThread):
    okSingal = pyqtSignal(int)
    progressSignal = pyqtSignal(int)

    def __init__(self, contact, messages):
        super().__init__()
        self.contact = contact
        self.messages = messages

    def run(self):
        origin_docx_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
        for message in self.messages:
            str_content = message[7]
            BytesExtra = message[10]
            timestamp = message[5]
            try:
                image_path = hard_link_db.get_image(str_content, BytesExtra, thumb=False)
                if not os.path.exists(os.path.join(MePC().wx_dir, image_path)):
                    image_thumb_path = hard_link_db.get_image(str_content, BytesExtra, thumb=True)
                    if not os.path.exists(os.path.join(MePC().wx_dir, image_thumb_path)):
                        continue
                    image_path = image_thumb_path
                image_path = get_image(image_path, base_path=f'/data/聊天记录/{self.contact.remark}/image')
                try:
                    os.utime(origin_docx_path + image_path[1:], (timestamp, timestamp))
                except:
                    pass
            except:
                logger.error(traceback.format_exc())
            finally:
                self.progressSignal.emit(1)
        self.okSingal.emit(47)
        print('图片子线程完成')


if __name__ == "__main__":
    from app.DataBase import micro_msg_db, misc_db
    from app.person import ContactPC
    from PyQt5.QtGui import QGuiApplication

    app = QGuiApplication([])
    contact_info_list = micro_msg_db.get_contact_by_username("wxid_lhbdvh3cnn4h22")
    contact_info = {
        'UserName': contact_info_list[0],
        'Alias': contact_info_list[1],
        'Type': contact_info_list[2],
        'Remark': contact_info_list[3],
        'NickName': contact_info_list[4],
        'smallHeadImgUrl': contact_info_list[7]
    }
    contact = ContactPC(contact_info)
    contact.smallHeadImgBLOG = misc_db.get_avatar_buffer(contact.wxid)
    contact.set_avatar(contact.smallHeadImgBLOG)
    mess = {1: True, 3: True, 34: True, 43: True, 47: True, 10000: True}
    MePC().name = "无题"
    MePC().wx_dir = r"C:\Users\HUAWEI\Documents\WeChat Files\wxid_05rvkbftizq822"
    MePC().wxid = "wxid_05rvkbftizq822"
    ChildThread(contact, 2, mess).to_html_()
    app.quit()
