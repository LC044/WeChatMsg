import hashlib
import os
import re
import threading
import time
import docx
import pandas as pd
import requests
from docx import shared
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from docxcompose.composer import Composer
import rcontact

import sys
from PyQt5.QtGui import QIcon
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import QPixmap

path = 'D:\Project\PythonProject\WeChat'
# conRemark = '曹雨萱'
# self_wxid = rcontact.get_self_wxid()
# ta_wxid = rcontact.get_one_wxid(conRemark)

'''
#! 创建emoji目录，存放emoji文件
'''


def mkdir(path):
    path = path.strip()
    path = path.rstrip("\\")
    if os.path.exists(path):
        return False
    os.makedirs(path)
    return True


mkdir(path+'/emoji')
# mkdir('..//db_tables')
'''
#! 将wxid使用MD5编码加密
#! 加密结果是用户头像路径
'''


def avatar_md5(wxid):
    m = hashlib.md5()
    # 参数必须是byte类型，否则报Unicode-objects must be encoded before hashing错误
    m.update(bytes(wxid.encode('utf-8')))
    return m.hexdigest()


'''
#! 获取头像文件完整路径
'''


def get_avator(wxid):
    avatar = avatar_md5(wxid)
    avatar_path = path + "/avatar/"
    Path = avatar_path + avatar[:2] + '/' + avatar[2:4]
    for root, dirs, files in os.walk(path):
        for file in files:
            if avatar in file:
                avatar = file
                break
    return Path + '/' + avatar

def read_csv(conRemark):
    '''
    :param conRemark: (str) 要导出的联系人备注名
    :return: pandas数据
    '''
    user_data = pd.read_csv(f'{path}/db_tables/{conRemark}.csv')
    '''将浮点数转化成字符串类型，否则会舍入影响时间结果'''
    user_data['createTime'] = user_data['createTime'].astype(str)
    # print(user_data)
    return user_data


def download_emoji(content, img_path):
    '''
    #! 下载emoji文件
    #!
    #!
    '''
    # if 1:
    try:
        # print(img_path)
        url = content.split('cdnurl = "')[1].split('"')[0]
        print(url)
        url = ':'.join(url.split('*#*'))
        if 'amp;' in url:
            url = ''.join(url.split('amp;'))
            print('emoji downloading!!!')
        resp = requests.get(url)
        with open(f'{path}/emoji/{img_path}', 'wb') as f:
            f.write(resp.content)
    except Exception:
        print("emoji download error")


def time_format(timestamp):
    '''
    #! 将字符串类型的时间戳转换成日期
    #! 返回格式化的时间字符串
    #! %Y-%m-%d %H:%M:%S
    '''
    # print(timestamp)
    # timestamp = timestamp[:-5]
    timestamp = float(timestamp[:-3] + '.' + timestamp[-3:])
    # print(timestamp)
    time_tuple = time.localtime(timestamp)
    # print(time.strftime("%Y-%m-%d %H:%M:%S", time_tuple))
    # quit()
    return time.strftime("%Y-%m-%d %H:%M:%S", time_tuple)





def IS_5_min(last_m, now_m):
    '''
    #! 判断两次聊天时间是不是大于五分钟
    #! 若大于五分钟则显示时间
    #! 否则不显示
    '''
    last_m = last_m[:-5]
    last_m = float(last_m + '.' + last_m[-5:2])
    now_m = now_m[:-5]
    now_m = float(now_m + '.' + now_m[-5:2])
    '''两次聊天记录时间差，单位是秒'''
    time_sub = now_m - last_m
    return time_sub >= 300


def judge_type(Type):
    pass







'''合并word文档到一个文件里'''


def merge_docx(conRemark, n):
    origin_docx_path = f"{path}/{conRemark}"
    all_word = os.listdir(origin_docx_path)
    all_file_path = []
    for i in range(n):
        file_name = f"{conRemark}{i}.docx"
        all_file_path.append(origin_docx_path + '/' + file_name)
    filename = f"{conRemark}.docx"
    # print(all_file_path)
    doc = docx.Document()
    doc.save(origin_docx_path + '/' + filename)
    master = docx.Document(origin_docx_path + '/' + filename)
    middle_new_docx = Composer(master)
    num = 0
    for word in all_file_path:
        word_document = docx.Document(word)
        word_document.add_page_break()
        if num != 0:
            middle_new_docx.append(word_document)
        num = num + 1
    middle_new_docx.save(origin_docx_path + '/' + filename)


class MyThread(QThread):
    signal = pyqtSignal(str)
    self_text = pyqtSignal(str)
    ta_text = pyqtSignal(str)
    bar = pyqtSignal(int)
    def __init__(self):
        super(MyThread, self).__init__()
        self.ta_info = {}
        self.self_info = {}
        self.textBrowser = None
        self.num = 0
        self.total_num = 1

    def get_avator(self):
        self.wxid = self.self_info['wxid']
        self.ta_wxid = self.ta_info['wxid']
        self.avator = get_avator(self.wxid)
        print(self.avator)
        self.ta_avator = get_avator(self.ta_wxid)
        print(self.ta_avator)
        # quit()
        self.img_self = open(self.avator, 'rb')
        self.img_ta = open(self.ta_avator, 'rb')

    def read_csv(self, conRemark):
        '''
        :param conRemark: (str) 要导出的联系人备注名
        :return: pandas数据
        '''
        user_data = pd.read_csv(f'{path}/db_tables/{conRemark}.csv')
        '''将浮点数转化成字符串类型，否则会舍入影响时间结果'''
        user_data['createTime'] = user_data['createTime'].astype(str)
        # print(user_data)
        self.total_num = len(user_data)
        return user_data
    def create_table(self,doc, isSend):
        '''
        #! 创建一个1*2表格
        #! isSend = 1 (0,0)存聊天内容，(0,1)存头像
        #! isSend = 0 (0,0)存头像，(0,1)存聊天内容
        #! 返回聊天内容的坐标
        '''
        table = doc.add_table(rows=1, cols=2, style='Normal Table')
        table.cell(0, 1).height = shared.Inches(0.5)
        table.cell(0, 0).height = shared.Inches(0.5)
        text_size = 1
        if isSend:
            '''表格右对齐'''
            table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            avatar = table.cell(0, 1).paragraphs[0].add_run()
            '''插入头像，设置头像宽度'''
            avatar.add_picture(self.img_self, width=shared.Inches(0.5))
            '''设置单元格宽度跟头像一致'''
            table.cell(0, 1).width = shared.Inches(0.5)
            content_cell = table.cell(0, 0)
            '''聊天内容右对齐'''
            content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            avatar = table.cell(0, 0).paragraphs[0].add_run()
            avatar.add_picture(self.img_ta, width=shared.Inches(0.5))
            '''设置单元格宽度'''
            table.cell(0, 0).width = shared.Inches(0.5)
            content_cell = table.cell(0, 1)
        '''聊天内容垂直居中对齐'''
        content_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        return content_cell

    def text(self, doc, isSend, message, status):
        if status == 5:
            message += '（未发出） '
        content_cell = self.create_table(doc, isSend)
        content_cell.paragraphs[0].add_run(message)
        content_cell.paragraphs[0].font_size = shared.Inches(0.5)
        self.self_text.emit(message)
        if isSend:
            p = content_cell.paragraphs[0]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph()

    def image(self, doc, isSend, Type, content, imgPath):
        '''
        #! 插入聊天图片
        #! isSend = 1 只有缩略图
        #! isSend = 0 有原图
        :param doc:
        :param isSend:
        :param Type:
        :param content:
        :param imgPath:
        :return:
        '''
        content = self.create_table(doc, isSend)
        run = content.paragraphs[0].add_run()
        imgPath = imgPath.split('//th_')[-1]
        Path = None
        if Type == 3:
            Path = f'{path}/image2//{imgPath[:2]}//{imgPath[2:4]}'
        elif Type == 47:
            Path = '{path}/emoji'
        for root, dirs, files in os.walk(Path):
            for file in files:
                if isSend:
                    if imgPath + 'hd' in file:
                        imgPath = file
                        try:
                            run.add_picture(f'{Path}/{imgPath}', height=shared.Inches(2))
                            doc.add_paragraph()
                        except Exception:
                            print("Error!image")
                        return
                elif imgPath in file:
                    imgPath = file
                    break
        try:
            run.add_picture(f'{Path}/{imgPath}', height=shared.Inches(2))
            doc.add_paragraph()
        except Exception:
            print("Error!image")

        # run.add_picture(f'{Path}/{imgPath}', height=shared.Inches(2))

    def emoji(self, doc, isSend, content, imgPath):
        '''
        #! 添加表情包
        :param isSend:
        :param content:
        :param imgPath:
        :return:
        '''
        if 1:
        # try:
            Path = f'{path}/emoji/{imgPath}'
            is_Exist = os.path.exists(Path)
            if not is_Exist:
                '''表情包不存在，则下载表情包到emoji文件夹中'''
                download_emoji(content, imgPath)
            self.image(doc, isSend, Type=47, content=content, imgPath=imgPath)
        # except Exception:
        #     print("can't find emoji!")

    def wx_file(self, doc, isSend, content, status):
        '''
        #! 添加微信文件
        :param isSend:
        :param content:
        :param status:
        :return:
        '''
        pattern = re.compile(r"<title>(.*?)<")
        r = pattern.search(content).group()
        filename = r.lstrip('<title>').rstrip('<')
        self.text(doc, isSend, filename, status)

    def retract_message(self,doc, isSend, content, status):
        '''
        #! 显示撤回消息
        :param isSend:
        :param content:
        :param status:
        :return:
        '''
        paragraph = doc.add_paragraph(content)
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def reply(self, doc, isSend, content, status):
        '''
        #! 添加回复信息
        :param isSend:
        :param content:
        :param status:
        :return:
        '''
        pattern1 = re.compile(r"<title>(?P<title>(.*?))</title>")
        title = pattern1.search(content).groupdict()['title']
        pattern2 = re.compile(r"<displayname>(?P<displayname>(.*?))</displayname>")
        displayname = pattern2.search(content).groupdict()['displayname']
        '''匹配回复的回复'''
        pattern3 = re.compile(r"\n?title&gt;(?P<content>(.*?))\n?&lt;/title&gt")
        if not pattern3.search(content):
            if isSend == 0:
                '''匹配对方的回复'''
                pattern3 = re.compile(r"<content>(?P<content>(.*?))</content>")
            else:
                '''匹配自己的回复'''
                pattern3 = re.compile(r"</msgsource>\n?<content>(?P<content>(.*?))\n?</content>")

        '''这部分代码完全可以用if代替'''

        try:
            '''试错'''
            text = pattern3.search(content).groupdict()['content']
        except Exception:
            try:
                '''试错'''
                text = pattern3.search(content).groupdict()['content']
            except Exception:
                '''试错'''
                pattern3 = re.compile(r"\n?<content>(?P<content>(.*?))\n?</content>")
                '''试错'''
                if pattern3.search(content):
                    text = pattern3.search(content).groupdict()['content']
                else:
                    text = '图片'
        if status == 5:
            message = '（未发出） ' + ''
        content_cell = self.create_table(doc, isSend)
        content_cell.paragraphs[0].add_run(title)
        content_cell.paragraphs[0].font_size = shared.Inches(0.5)
        reply_p = content_cell.add_paragraph()
        run = content_cell.paragraphs[1].add_run(displayname + ':' + text)
        '''设置被回复内容格式'''
        run.font.color.rgb = shared.RGBColor(121, 121, 121)
        run.font_size = shared.Inches(0.3)
        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25

        if isSend:
            p = content_cell.paragraphs[0]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            reply_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph()


    def pat_a_pat(self, doc, isSend, content, status):
        '''
        #! 添加拍一拍信息
        todo 把wxid转化成昵称
        :param isSend:
        :param content:
        :param status:
        :return:
        '''
        '''<template><![CDATA[我${fromusername@textstatusicon}拍了拍自己]]></template>'''
        pattern = re.compile(r"<template><!\[CDATA\[(?P<it>(.*?))]]></template>")
        result = pattern.search(content).groupdict()['it']
        fromusername = '${fromusername@textstatusicon}'
        pattedusername = '${pattedusername@textstatusicon}'
        '''我拍别人'''
        if result[0] == u'我':
            result = ''.join(result.split(fromusername))
            result = ''.join(result.split(pattedusername))
            pat = result
        else:
            '''处理多余的引号'''
            result = result.split('""') if '""' in result else result.split('"')
            for i in range(len(result)):
                if fromusername in result[i]:
                    result[i] = result[i].rstrip(fromusername)
                elif pattedusername in result[i]:
                    result[i] = result[i].rstrip(pattedusername)

            if len(result) >= 4:
                '''别人拍别人
                #! ""${s407575157}${fromusername@textstatusicon}"" \
                #! 拍了拍 \
                #! ""${wxid_7rs401fwlaje22}${pattedusername@textstatusicon}"" \
                #! 的豪宅不小心塌了??
                #! [' ', wxid0, '拍了拍', wxid1, '内容']
                '''
                wxid0 = result[1].lstrip('${').rstrip('}')  # ! 第一个人的wxid
                wxid1 = result[3].lstrip('${').rstrip('}')  # ! 第二个人的wxid
                nickname0 = rcontact.wxid_to_conRemark(wxid0)  # ! 将wxid转换成昵称
                nickname1 = rcontact.wxid_to_conRemark(wxid1)  # ! 将wxid转换成昵称
                pat = nickname0 + result[2] + nickname1  # todo 留着把wxid转换成昵称
                if len(result) == 5:
                    pat += result[4]
            else:
                '''#!  ""${wxid_8piw6sb4hvfm22}"" 拍了拍我  '''
                '''
                #! 别人拍我
                #! [' ', wxid0, '拍了拍我']
                '''
                wxid0 = result[1].lstrip('${').rstrip('}')
                pat = wxid0 + result[2]
        print(pat)
        p = doc.add_paragraph()
        run = p.add_run(pat)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        '''设置拍一拍文字格式'''
        run.font.color.rgb = shared.RGBColor(121, 121, 121)
        run.font_size = shared.Inches(0.3)
        # run.font.highlight_color=WD_COLOR_INDEX.GRAY_25

    def video(self, doc, isSend, content, status, img_path):
        print(content, img_path)

    def to_docx(self, user_data, i, conRemark):
        '''

        :param user_data:
        :param i:
        :param conRemark:
        :return:
        '''
        '''创建联系人目录'''
        mkdir(f"{path}/{conRemark}")
        filename = f"{path}/{conRemark}/{conRemark}{i}.docx"
        doc = docx.Document()
        now_timestamp = '1600008700000.0'
        for row_index, row in user_data.iterrows():
            self.num += 1
            self.bar.emit(int((self.num)/self.total_num*100))
            Type = row['type']
            content = row['content']
            isSend = row['isSend']
            last_timestamp = now_timestamp
            now_timestamp = row['createTime']
            createTime = time_format(now_timestamp)
            imgPath = row['imgPath']
            status = row['status']
            if IS_5_min(last_timestamp, now_timestamp):
                doc.add_paragraph(createTime).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if Type == 1:
                # print('文本信息')
                # print(createTime, content)
                # continue
                self.text(doc, isSend, content, status)
            elif Type == 3:
                # print('图片信息')
                # continue
                self.image(doc, isSend, 3, content, imgPath)
            elif Type == 47:
                # print(content)
                print(imgPath,content)
                self.emoji(doc, isSend, content, imgPath)
            elif Type == 1090519089:
                self.wx_file(doc, isSend, content, status)
            elif Type == 268445456:
                self.retract_message(doc, isSend, content, status)
            elif Type == 822083633:
                self.reply(doc, isSend, content, status)
            elif Type == 922746929:
                self.pat_a_pat(doc, isSend, content, status)
            elif Type == 43:
                # print(createTime)
                self.video(doc, isSend, content, status, imgPath)
        # doc.add_paragraph(str(i))
        # print(filename)
        doc.save(filename)
    def run(self):
        if 1:
        # try:
            self.get_avator()
            conRemark = self.ta_info['conRemark']
            self.self_text.emit(conRemark)
            self.self_text.emit(path)
            user_data = self.read_csv(conRemark)
            l = len(user_data)
            n = 50
            threads = []
            for i in range(n):
                q = i * (l // n)
                p = (i + 1) * (l // n)
                if i == n - 1:
                    p = l
                data = user_data[q:p]
                self.to_docx(data, i, conRemark)
            # #     t = threading.Thread(target=self.to_docx, args=(data, i, conRemark))
            # #     threads.append(t)
            # # for thr in threads:
            # #     thr.start()
            # #     thr.join()
            self.self_text.emit('\n\n\n导出进度还差一点点！！！')
            self.bar.emit(99)
            merge_docx(conRemark, n)
            self.self_text.emit(f'{conRemark}聊天记录导出成功！！！')
            self.bar.emit(100)
        # except Exception as e:
        #     self.self_text.emit('发生异常')
        #     print(e)
            #self.self_text.emit(e)
if __name__ == '__main__':
    # # conRemark = '张三' #! 微信备注名
    # n = 100  # ! 分割的文件个数
    # main(conRemark, n)
    # img_self.close()
    # img_ta.close()
    t = MyThread()
    # t.ta_info = {
    #     'wxid': 'wxid_q3ozn70pweud22',
    #     'conRemark': '小钱'
    # }
    t.ta_info = {
        'wxid': 'wxid_8piw6sb4hvfm22',
        'conRemark': '曹雨萱'
    }
    #wxid_8piw6sb4hvfm22
    t.self_info = {
        'wxid': 'wxid_27hqbq7vx5hf22',
        'conRemark': 'Shuaikang Zhou'
    }
    t.run()
