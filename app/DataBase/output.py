import csv
import os
import time
import traceback
from typing import List

import docx
from PyQt5.QtCore import pyqtSignal, QThread, QObject
from PyQt5.QtWidgets import QFileDialog
from docx.oxml.ns import qn
from docxcompose.composer import Composer

from app.DataBase.exporter_csv import CSVExporter
from app.DataBase.exporter_docx import DocxExporter
from app.DataBase.exporter_html import HtmlExporter
from app.DataBase.exporter_txt import TxtExporter
from app.DataBase.hard_link import decodeExtraBuf
from app.config import output_dir
from .package_msg import PackageMsg
from ..DataBase import media_msg_db, hard_link_db, micro_msg_db, msg_db
from ..log import logger
from ..person import Me
from ..util.image import get_image

os.makedirs(os.path.join(output_dir, '聊天记录'), exist_ok=True)


class Output(QThread):
    """
    发送信息线程
    """
    startSignal = pyqtSignal(int)
    progressSignal = pyqtSignal(int)
    rangeSignal = pyqtSignal(int)
    okSignal = pyqtSignal(int)
    batchOkSignal = pyqtSignal(int)
    nowContact = pyqtSignal(str)
    i = 1
    CSV = 0
    DOCX = 1
    HTML = 2
    CSV_ALL = 3
    CONTACT_CSV = 4
    TXT = 5
    Batch = 10086

    def __init__(self, contact, type_=DOCX, message_types={}, sub_type=[], time_range=None, parent=None):
        super().__init__(parent)
        self.children = []
        self.last_timestamp = 0
        self.sub_type = sub_type
        self.time_range = time_range
        self.message_types = message_types
        self.sec = 2  # 默认1000秒
        self.contact = contact
        self.msg_id = 0
        self.output_type: int | List[int] = type_
        self.total_num = 1
        self.num = 0

    def progress(self, value):
        self.progressSignal.emit(value)

    def output_image(self):
        """
        导出全部图片
        @return:
        """
        return

    def output_emoji(self):
        """
        导出全部表情包
        @return:
        """
        return

    def to_csv_all(self):
        """
        导出全部聊天记录到CSV
        @return:
        """

        origin_path = os.path.join(os.path.abspath('.'), output_dir, '聊天记录')
        os.makedirs(origin_path, exist_ok=True)
        filename = QFileDialog.getSaveFileName(None, "save file", os.path.join(os.getcwd(), 'messages.csv'),
                                               "csv files (*.csv);;all files(*.*)")
        if not filename[0]:
            return
        self.startSignal.emit(1)
        filename = filename[0]
        # columns = ["用户名", "消息内容", "发送时间", "发送状态", "消息类型", "isSend", "msgId"]
        columns = ['localId', 'TalkerId', 'Type', 'SubType',
                   'IsSender', 'CreateTime', 'Status', 'StrContent',
                   'StrTime', 'Remark', 'NickName', 'Sender']

        packagemsg = PackageMsg()
        messages = packagemsg.get_package_message_all()
        # 写入CSV文件
        with open(filename, mode='w', newline='', encoding='utf-8-sig') as file:
            writer = csv.writer(file)
            writer.writerow(columns)
            # 写入数据
            writer.writerows(messages)
        self.okSignal.emit(1)

    def contact_to_csv(self):
        """
        导出联系人到CSV
        @return:
        """
        filename = QFileDialog.getSaveFileName(None, "save file", os.path.join(os.getcwd(), 'contacts.csv'),
                                               "csv files (*.csv);;all files(*.*)")
        if not filename[0]:
            return
        self.startSignal.emit(1)
        filename = filename[0]
        # columns = ["用户名", "消息内容", "发送时间", "发送状态", "消息类型", "isSend", "msgId"]
        columns = ['UserName', 'Alias', 'Type', 'Remark', 'NickName', 'PYInitial', 'RemarkPYInitial', 'smallHeadImgUrl',
                   'bigHeadImgUrl', 'label', 'gender', 'telephone', 'signature', 'country/region', 'province', 'city']
        contacts = micro_msg_db.get_contact()
        # 写入CSV文件
        with open(filename, mode='w', newline='', encoding='utf-8-sig') as file:
            writer = csv.writer(file)
            writer.writerow(columns)
            # 写入数据
            # writer.writerows(contacts)
            for contact in contacts:
                detail = decodeExtraBuf(contact[9])
                gender_code = detail.get('gender')
                if gender_code == 0:
                    gender = '未知'
                elif gender_code == 1:
                    gender = '男'
                else:
                    gender = '女'
                writer.writerow([*contact[:9], contact[10], gender, detail.get('telephone'), detail.get('signature'),
                                 *detail.get('region')])

        self.okSignal.emit(1)

    def batch_export(self):
        print('开始批量导出')
        print(self.sub_type, self.message_types)
        print(len(self.contact))
        print([contact.remark for contact in self.contact])
        self.batch_num_total = len(self.contact) * len(self.sub_type)
        self.batch_num = 0
        self.rangeSignal.emit(self.batch_num_total)
        for contact in self.contact:
            # print('联系人', contact.remark)
            for type_ in self.sub_type:
                # print('导出类型', type_)
                if type_ == self.DOCX:
                    self.to_docx(contact, self.message_types, True)
                elif type_ == self.TXT:
                    # print('批量导出txt')
                    self.to_txt(contact, self.message_types, True)
                elif type_ == self.CSV:
                    self.to_csv(contact, self.message_types, True)
                elif type_ == self.HTML:
                    self.to_html(contact, self.message_types, True)

    def batch_finish_one(self, num):
        self.nowContact.emit(self.contact[self.batch_num // len(self.sub_type)].remark)
        self.batch_num += 1
        if self.batch_num == self.batch_num_total:
            self.okSignal.emit(1)

    def merge_docx(self, n):
        conRemark = self.contact.remark
        origin_path = os.path.join(os.path.abspath('.'), output_dir, '聊天记录',conRemark)
        filename = f"{origin_path}/{conRemark}_{n}.docx"
        if n == 10086:
            # self.document.append(self.document)
            file = os.path.join(origin_path, f'{conRemark}.docx')
            try:
                self.document.save(file)
            except PermissionError:
                file = file[:-5] + f'{time.time()}' + '.docx'
                self.document.save(file)
            self.okSignal.emit(1)
            return
        doc = docx.Document(filename)
        self.document.append(doc)
        os.remove(filename)
        if n % 50 == 0:
            # self.document.append(self.document)
            file = os.path.join(origin_path, f'{conRemark}-{n // 50}.docx')
            try:
                self.document.save(file)
            except PermissionError:
                file = file[:-5] + f'{time.time()}' + '.docx'
                self.document.save(file)
            doc = docx.Document()
            doc.styles["Normal"].font.name = "Cambria"
            doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            self.document = Composer(doc)

    def to_docx(self, contact, message_types, is_batch=False):
        doc = docx.Document()
        doc.styles["Normal"].font.name = "Cambria"
        doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        self.document = Composer(doc)
        Child = DocxExporter(contact, type_=self.DOCX, message_types=message_types, time_range=self.time_range)
        self.children.append(Child)
        Child.progressSignal.connect(self.progress)
        if not is_batch:
            Child.rangeSignal.connect(self.rangeSignal)
        Child.okSignal.connect(self.merge_docx if not is_batch else self.batch_finish_one)
        Child.start()

    def to_txt(self, contact, message_types, is_batch=False):
        Child = TxtExporter(contact, type_=self.TXT, message_types=message_types, time_range=self.time_range)
        self.children.append(Child)
        Child.progressSignal.connect(self.progress)
        if not is_batch:
            Child.rangeSignal.connect(self.rangeSignal)
        Child.okSignal.connect(self.okSignal if not is_batch else self.batch_finish_one)
        Child.start()

    def to_html(self, contact, message_types, is_batch=False):
        Child = HtmlExporter(contact, type_=self.output_type, message_types=message_types, time_range=self.time_range)
        self.children.append(Child)
        Child.progressSignal.connect(self.progress)
        if not is_batch:
            Child.rangeSignal.connect(self.rangeSignal)
        Child.okSignal.connect(self.count_finish_num)
        Child.start()
        self.total_num = 1
        if message_types.get(34):
            # 语音消息单独的线程
            self.total_num += 1
            output_media = OutputMedia(contact, time_range=self.time_range)
            self.children.append(output_media)
            output_media.okSingal.connect(self.count_finish_num)
            output_media.progressSignal.connect(self.progressSignal)
            output_media.start()
        if message_types.get(47):
            # emoji消息单独的线程
            self.total_num += 1
            output_emoji = OutputEmoji(contact, time_range=self.time_range)
            self.children.append(output_emoji)
            output_emoji.okSingal.connect(self.count_finish_num)
            output_emoji.progressSignal.connect(self.progressSignal)
            output_emoji.start()
        if message_types.get(3):
            # 图片消息单独的线程
            self.total_num += 1
            output_image = OutputImage(contact, time_range=self.time_range)
            self.children.append(output_image)
            output_image.okSingal.connect(self.count_finish_num)
            output_image.progressSignal.connect(self.progressSignal)
            output_image.start()

    def to_csv(self, contact, message_types, is_batch=False):
        Child = CSVExporter(contact, type_=self.CSV, message_types=message_types, time_range=self.time_range)
        self.children.append(Child)
        Child.progressSignal.connect(self.progress)
        if not is_batch:
            Child.rangeSignal.connect(self.rangeSignal)
        Child.okSignal.connect(self.okSignal if not is_batch else self.batch_finish_one)
        Child.start()

    def run(self):
        if self.output_type == self.DOCX:
            self.to_docx(self.contact, self.message_types)
        elif self.output_type == self.CSV_ALL:
            self.to_csv_all()
        elif self.output_type == self.CONTACT_CSV:
            self.contact_to_csv()
        elif self.output_type == self.TXT:
            self.to_txt(self.contact, self.message_types)
        elif self.output_type == self.CSV:
            self.to_csv(self.contact, self.message_types)
        elif self.output_type == self.HTML:
            self.to_html(self.contact, self.message_types)
        elif self.output_type == self.Batch:
            self.batch_export()

    def count_finish_num(self, num):
        """
        记录子线程完成个数
        @param num:
        @return:
        """
        self.num += 1
        if self.num == self.total_num:
            # 所有子线程都完成之后就发送完成信号
            if self.output_type == self.Batch:
                self.batch_finish_one(1)
            else:
                self.okSignal.emit(1)
            self.num = 0

    def cancel(self):
        self.requestInterruption()


class OutputMedia(QThread):
    """
    导出语音消息
    """
    okSingal = pyqtSignal(int)
    progressSignal = pyqtSignal(int)

    def __init__(self, contact, time_range=None):
        super().__init__()
        self.contact = contact
        self.time_range = time_range

    def run(self):
        origin_path = os.path.join(os.path.abspath('.'),output_dir,'聊天记录',self.contact.remark)
        messages = msg_db.get_messages_by_type(self.contact.wxid, 34, time_range=self.time_range)
        for message in messages:
            is_send = message[4]
            msgSvrId = message[9]
            try:
                audio_path = media_msg_db.get_audio(msgSvrId, output_path=origin_path + "/voice")
            except:
                logger.error(traceback.format_exc())
            finally:
                self.progressSignal.emit(1)
        self.okSingal.emit(34)


class OutputEmoji(QThread):
    """
    导出表情包
    """
    okSingal = pyqtSignal(int)
    progressSignal = pyqtSignal(int)

    def __init__(self, contact, time_range=None):
        super().__init__()
        self.contact = contact
        self.time_range = time_range

    def run(self):
        origin_path = os.path.join(os.path.abspath('.'),output_dir,'聊天记录',self.contact.remark)
        messages = msg_db.get_messages_by_type(self.contact.wxid, 47, time_range=self.time_range)
        for message in messages:
            str_content = message[7]
            try:
                pass
                # emoji_path = get_emoji(str_content, thumb=True, output_path=origin_path + '/emoji')
            except:
                logger.error(traceback.format_exc())
            finally:
                self.progressSignal.emit(1)
        self.okSingal.emit(47)


class OutputImage(QThread):
    """
    导出图片
    """
    okSingal = pyqtSignal(int)
    progressSignal = pyqtSignal(int)

    def __init__(self, contact, time_range):
        super().__init__()
        self.contact = contact
        self.child_thread_num = 2
        self.time_range = time_range
        self.child_threads = [0] * (self.child_thread_num + 1)
        self.num = 0

    def count1(self, num):
        self.num += 1
        print('图片导出完成一个')
        if self.num == self.child_thread_num:
            self.okSingal.emit(47)
            print('图片导出完成')

    def run(self):
        origin_path = os.path.join(os.path.abspath('.'),output_dir,'聊天记录',self.contact.remark)
        messages = msg_db.get_messages_by_type(self.contact.wxid, 3, time_range=self.time_range)
        base_path = os.path.join(output_dir,'聊天记录',self.contact.remark,'image')
        for message in messages:
            str_content = message[7]
            BytesExtra = message[10]
            timestamp = message[5]
            try:
                image_path = hard_link_db.get_image(str_content, BytesExtra, up_dir=Me().wx_dir, thumb=False)
                image_path = get_image(image_path, base_path=base_path)
                try:
                    os.utime(origin_path + image_path[1:], (timestamp, timestamp))
                except:
                    pass
            except:
                logger.error(traceback.format_exc())
            finally:
                self.progressSignal.emit(1)
        self.okSingal.emit(47)


class OutputImageChild(QThread):
    okSingal = pyqtSignal(int)
    progressSignal = pyqtSignal(int)

    def __init__(self, contact, messages, time_range):
        super().__init__()
        self.contact = contact
        self.messages = messages
        self.time_range = time_range

    def run(self):
        origin_path = os.path.join(os.path.abspath('.'),output_dir,'聊天记录',self.contact.remark)
        for message in self.messages:
            str_content = message[7]
            BytesExtra = message[10]
            timestamp = message[5]
            try:
                image_path = hard_link_db.get_image(str_content, BytesExtra, thumb=False)
                if not os.path.exists(os.path.join(Me().wx_dir, image_path)):
                    image_thumb_path = hard_link_db.get_image(str_content, BytesExtra, thumb=True)
                    if not os.path.exists(os.path.join(Me().wx_dir, image_thumb_path)):
                        continue
                    image_path = image_thumb_path
                image_path = get_image(image_path, base_path=f'/data/聊天记录/{self.contact.remark}/image')
                try:
                    os.utime(origin_path + image_path[1:], (timestamp, timestamp))
                except:
                    pass
            except:
                logger.error(traceback.format_exc())
            finally:
                self.progressSignal.emit(1)
        self.okSingal.emit(47)
        print('图片子线程完成')


if __name__ == "__main__":
    pass
