import os
import shutil
import time
from re import findall

import docx
import unicodedata
from docx import shared
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docxcompose.composer import Composer

from app.DataBase import msg_db, hard_link_db
from app.DataBase.exporter import ExporterBase, escape_js_and_html
from app.config import output_dir
from app.log import logger
from app.person import Me
from app.util.compress_content import parser_reply, share_card, music_share
from app.util.image import get_image_abs_path
from app.util.music import get_music_path

# 要删除的编码字符
encoded_chars = b'\x00\x01\x02\x03\x04\x05\x06\x07\x08\x0b\x0c\x0e\x0f\x10\x11\x12\x13\x14\x15\x16\x17\x18\x19\x1a\x1b\x1c\x1d\x1e\x1f'

# 创建一个字典，将要删除的字符映射为 None
char_mapping = {char: None for char in encoded_chars}


def filter_control_characters(input_string):
    """
    过滤掉不可打印字符
    @param input_string:
    @return:
    """

    # 过滤掉非可打印字符
    filtered_string = input_string.translate(char_mapping)

    return filtered_string


class DocxExporter(ExporterBase):
    def text(self, doc, message):
        type_ = message[2]
        str_content = message[7]
        str_time = message[8]
        is_send = message[4]
        timestamp = message[5]
        is_chatroom = 1 if self.contact.is_chatroom else 0

        display_name = self.get_display_name(is_send, message)
        avatar = self.get_avatar_path(is_send, message, True)
        content_cell = self.create_table(doc, is_send, avatar)
        try:
            content_cell.paragraphs[0].add_run(str_content)
        except ValueError:
            try:
                str_content = filter_control_characters(str_content)
                content_cell.paragraphs[0].add_run(str_content)
            except ValueError:
                logger.error(f'非法字符:{str_content}')
                content_cell.paragraphs[0].add_run('非法字符')
        content_cell.paragraphs[0].font_size = shared.Inches(0.5)
        # doc.add_picture(avatar)
        if is_send:
            p = content_cell.paragraphs[0]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph()

    def image(self, doc, message):
        str_content = message[7]
        is_send = message[4]
        BytesExtra = message[10]
        avatar = self.get_avatar_path(is_send, message, True)
        content = self.create_table(doc, is_send, avatar)
        run = content.paragraphs[0].add_run()
        str_content = escape_js_and_html(str_content)
        image_path = hard_link_db.get_image(str_content, BytesExtra, thumb=True)
        base_path = os.path.join(output_dir, '聊天记录', self.contact.remark, 'image')
        if not os.path.exists(os.path.join(Me().wx_dir, image_path)):
            image_thumb_path = hard_link_db.get_image(str_content, BytesExtra, thumb=False)
            if not os.path.exists(os.path.join(Me().wx_dir, image_thumb_path)):
                return
            image_path = image_thumb_path
        image_path = get_image_abs_path(image_path, base_path=base_path)
        try:
            run.add_picture(image_path, height=shared.Inches(2))
            doc.add_paragraph()
        except Exception:
            print("Error!image")

    def audio(self, doc, message):
        str_content = message[7]
        str_time = message[8]
        is_send = message[4]
        msgSvrId = message[9]
        timestamp = message[5]
        is_chatroom = 1 if self.contact.is_chatroom else 0
        avatar = self.get_avatar_path(is_send, message)
        display_name = self.get_display_name(is_send, message)
        avatar = self.get_avatar_path(is_send, message, True)
        content_cell = self.create_table(doc, is_send, avatar)
        content_cell.paragraphs[0].add_run('【语音】')
        content_cell.paragraphs[0].font_size = shared.Inches(0.5)
        if is_send:
            p = content_cell.paragraphs[0]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph()

    def emoji(self, doc, message):
        str_content = message[7]
        str_time = message[8]
        is_send = message[4]
        timestamp = message[5]
        is_chatroom = 1 if self.contact.is_chatroom else 0
        avatar = self.get_avatar_path(is_send, message)
        display_name = self.get_display_name(is_send, message)
        avatar = self.get_avatar_path(is_send, message, True)
        content_cell = self.create_table(doc, is_send, avatar)
        content_cell.paragraphs[0].add_run('【表情包】')
        content_cell.paragraphs[0].font_size = shared.Inches(0.5)
        if is_send:
            p = content_cell.paragraphs[0]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph()

    def file(self, doc, message):
        bytesExtra = message[10]
        str_time = message[8]
        is_send = message[4]
        timestamp = message[5]
        is_chatroom = 1 if self.contact.is_chatroom else 0
        avatar = self.get_avatar_path(is_send, message)
        display_name = self.get_display_name(is_send, message)
        avatar = self.get_avatar_path(is_send, message, True)
        content_cell = self.create_table(doc, is_send, avatar)
        content_cell.paragraphs[0].add_run('【文件】')
        content_cell.paragraphs[0].font_size = shared.Inches(0.5)
        if is_send:
            p = content_cell.paragraphs[0]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph()

    def refermsg(self, doc, message):
        """
        处理回复消息
        @param doc:
        @param message:
        @return:
        """
        str_time = message[8]
        is_send = message[4]
        content = parser_reply(message[11])
        refer_msg = content.get('refer')
        timestamp = message[5]
        is_chatroom = 1 if self.contact.is_chatroom else 0
        avatar = self.get_avatar_path(is_send, message)
        display_name = self.get_display_name(is_send, message)
        avatar = self.get_avatar_path(is_send, message, True)
        content_cell = self.create_table(doc, is_send, avatar)
        content_cell.paragraphs[0].add_run(content.get('title'))
        content_cell.paragraphs[0].font_size = shared.Inches(0.5)
        reply_p = content_cell.add_paragraph()
        reply_content = f"{refer_msg.get('displayname')}:{refer_msg.get('content')}" if refer_msg else '未知引用'
        run = content_cell.paragraphs[1].add_run(reply_content)
        '''设置被回复内容格式'''
        run.font.color.rgb = shared.RGBColor(121, 121, 121)
        run.font_size = shared.Inches(0.3)
        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25

        if is_send:
            p = content_cell.paragraphs[0]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            reply_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph()

    def system_msg(self, doc, message):
        str_content = message[7]
        is_send = message[4]
        str_time = message[8]
        timestamp = message[5]
        is_chatroom = 1 if self.contact.is_chatroom else 0

        str_content = str_content.replace('<![CDATA[', "").replace(
            ' <a href="weixin://revoke_edit_click">重新编辑</a>]]>', "")
        res = findall('(</{0,1}(img|revo|_wc_cus|a).*?>)', str_content)
        for xmlstr, b in res:
            str_content = str_content.replace(xmlstr, "")
        doc.add_paragraph(str_content).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def video(self, doc, message):
        type_ = message[2]
        str_content = message[7]
        str_time = message[8]
        is_send = message[4]
        BytesExtra = message[10]
        timestamp = message[5]
        is_chatroom = 1 if self.contact.is_chatroom else 0
        avatar = self.get_avatar_path(is_send, message)
        display_name = self.get_display_name(is_send, message)
        avatar = self.get_avatar_path(is_send, message, True)
        content_cell = self.create_table(doc, is_send, avatar)
        content_cell.paragraphs[0].add_run('【视频】')
        content_cell.paragraphs[0].font_size = shared.Inches(0.5)
        if is_send:
            p = content_cell.paragraphs[0]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph()

    def create_table(self, doc, is_send, avatar_path):
        '''
        #! 创建一个1*2表格
        #! isSend = 1 (0,0)存聊天内容，(0,1)存头像
        #! isSend = 0 (0,0)存头像，(0,1)存聊天内容
        #! 返回聊天内容的坐标
        '''
        table = doc.add_table(rows=1, cols=2, style='Normal Table')
        table.cell(0, 1).height = shared.Inches(0.5)
        table.cell(0, 0).height = shared.Inches(0.5)
        if is_send:
            '''表格右对齐'''
            table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            avatar = table.cell(0, 1).paragraphs[0].add_run()
            '''插入头像，设置头像宽度'''
            avatar.add_picture(avatar_path, width=shared.Inches(0.5))
            '''设置单元格宽度跟头像一致'''
            table.cell(0, 1).width = shared.Inches(0.5)
            content_cell = table.cell(0, 0)
            '''聊天内容右对齐'''
            content_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            avatar = table.cell(0, 0).paragraphs[0].add_run()
            avatar.add_picture(avatar_path, width=shared.Inches(0.5))
            '''设置单元格宽度'''
            table.cell(0, 0).width = shared.Inches(0.5)
            content_cell = table.cell(0, 1)
        '''聊天内容垂直居中对齐'''
        content_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        return content_cell

    def music_share(self, doc, message):
        origin_path = os.path.join(os.path.abspath('.'), output_dir, '聊天记录', self.contact.remark)
        is_send = message[4]
        timestamp = message[5]
        content = music_share(message[11])
        music_path = ''
        if content.get('audio_url') != '':
            music_path = get_music_path(content.get('audio_url'), content.get('title'),
                                        output_path=origin_path + '/music')
            if music_path != '':
                music_path = f'./music/{os.path.basename(music_path)}'
                music_path = music_path.replace('\\', '/')
        is_chatroom = 1 if self.contact.is_chatroom else 0
        avatar = self.get_avatar_path(is_send, message)
        display_name = self.get_display_name(is_send, message)

    def share_card(self, doc, message):
        origin_path = f"{os.path.abspath('.')}/data/聊天记录/{self.contact.remark}"
        is_send = message[4]
        timestamp = message[5]
        bytesExtra = message[10]
        compress_content_ = message[11]
        card_data = share_card(bytesExtra, compress_content_)
        is_chatroom = 1 if self.contact.is_chatroom else 0
        avatar = self.get_avatar_path(is_send, message)
        display_name = self.get_display_name(is_send, message)
        thumbnail = ''
        if card_data.get('thumbnail'):
            thumbnail = os.path.join(Me().wx_dir, card_data.get('thumbnail'))
            if os.path.exists(thumbnail):
                shutil.copy(thumbnail, os.path.join(origin_path, 'image', os.path.basename(thumbnail)))
                thumbnail = './image/' + os.path.basename(thumbnail)
            else:
                thumbnail = ''
        app_logo = ''
        if card_data.get('app_logo'):
            app_logo = os.path.join(Me().wx_dir, card_data.get('app_logo'))
            if os.path.exists(app_logo):
                shutil.copy(app_logo, os.path.join(origin_path, 'image', os.path.basename(app_logo)))
                app_logo = './image/' + os.path.basename(app_logo)
            else:
                app_logo = ''

    def merge_docx(self, conRemark, n):
        origin_path = os.path.join(os.path.abspath('.'), output_dir, '聊天记录')
        all_file_path = []
        for i in range(n):
            file_name = f"{conRemark}{i}.docx"
            all_file_path.append(origin_path + '/' + file_name)
        filename = f"{conRemark}.docx"
        # print(all_file_path)
        doc = docx.Document()
        doc.save(origin_path + '/' + filename)
        master = docx.Document(origin_path + '/' + filename)
        middle_new_docx = Composer(master)
        num = 0
        for word in all_file_path:
            word_document = docx.Document(word)
            word_document.add_page_break()
            if num != 0:
                middle_new_docx.append(word_document)
            num = num + 1
            os.remove(word)
        middle_new_docx.save(origin_path + '/' + filename)

    def export(self):
        print(f"【开始导出 DOCX {self.contact.remark}】")
        origin_path = os.path.join(os.path.abspath('.'), output_dir, '聊天记录', self.contact.remark)
        messages = msg_db.get_messages(self.contact.wxid, time_range=self.time_range)
        Me().save_avatar(os.path.join(origin_path, 'avatar', f'{Me().wxid}.png'))
        if self.contact.is_chatroom:
            for message in messages:
                if message[4]:  # is_send
                    continue
                try:
                    chatroom_avatar_path =os.path.join(origin_path, 'avatar', f'{message[13].wxid}.png')
                    message[13].save_avatar(chatroom_avatar_path)
                except:
                    print(message)
                    pass
        else:
            self.contact.save_avatar(os.path.join(origin_path, 'avatar', f'{self.contact.wxid}.png'))
        self.rangeSignal.emit(len(messages))

        def newdoc():
            nonlocal n, doc
            doc = docx.Document()
            doc.styles["Normal"].font.name = "Cambria"
            doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            n += 1

        doc = None
        n = 0
        index = 0
        newdoc()
        for index, message in enumerate(messages):
            if index % 200 == 0 and index:
                filename = os.path.join(origin_path, f"{self.contact.remark}_{n}.docx")
                doc.save(filename)
                self.okSignal.emit(n)
                newdoc()

            type_ = message[2]
            sub_type = message[3]
            timestamp = message[5]
            self.progressSignal.emit(1)
            if self.is_5_min(timestamp):
                str_time = message[8]
                doc.add_paragraph(str_time).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if type_ == 1 and self.message_types.get(type_):
                self.text(doc, message)
            elif type_ == 3 and self.message_types.get(type_):
                self.image(doc, message)
            elif type_ == 34 and self.message_types.get(type_):
                self.audio(doc, message)
            elif type_ == 43 and self.message_types.get(type_):
                self.video(doc, message)
            elif type_ == 47 and self.message_types.get(type_):
                self.emoji(doc, message)
            elif type_ == 10000 and self.message_types.get(type_):
                self.system_msg(doc, message)
            elif type_ == 49 and sub_type == 57 and self.message_types.get(1):
                self.refermsg(doc, message)
            elif type_ == 49 and sub_type == 6 and self.message_types.get(4906):
                self.file(doc, message)
            if index % 25 == 0:
                print(f"【导出 DOCX {self.contact.remark}】{index}/{len(messages)}")
        if index % 25:
            print(f"【导出 DOCX {self.contact.remark}】{index + 1}/{len(messages)}")
        filename = os.path.join(origin_path, f"{self.contact.remark}_{n}.docx")
        try:
            # document.save(filename)
            doc.save(filename)
        except PermissionError:
            filename = filename[:-5] + f'{time.time()}' + '.docx'
            # document.save(filename)
            doc.save(filename)
        self.okSignal.emit(n)
        print(f"【完成导出 DOCX {self.contact.remark}】")
        self.okSignal.emit(10086)
